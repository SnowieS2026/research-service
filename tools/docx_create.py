"""
docx_create.py -- Create Word (.docx) documents from structured input.
Usage:
  python docx_create.py --output "out.docx" --title "Title" --content "Block1|Block2|..."
  python docx_create.py --output "out.docx" --file "source.md"

Blocks in --content:
  Heading 1|text      -> Heading 1 style
  Heading 2|text      -> Heading 2 style
  Heading 3|text      -> Heading 3 style
  Bold|text           -> Bold paragraph
  Bullet|item1|item2  -> Bulleted list
  Number|item1|item2  -> Numbered list
  Roman|item1|item2   -> Roman numeral list
  Quote|text          -> Block quote
  HR                  -> Horizontal rule
  |text               -> Normal paragraph

Inline formatting: **bold**, *italic*, __underline__

Options:
  --output             Output .docx path (required)
  --title              Document title
  --author             Author name
  --content            Pipe-delimited content blocks
  --file               Load content from a markdown file instead
  --table              Table data: "Header1,Header2|Row1C1,Row1C2|Row2C1,Row2C2"
  --image              Image: "path,width,alignment" (width in px, alignment: left/center/right)
  --hyperlink          Hyperlink: "link text|url"
  --stylesheet         Use default Word styles (Heading 1, etc.) instead of direct formatting
  --no-inline          Disable inline markdown formatting
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def parse_inline(text):
    """Parse inline **bold**, *italic*, __underline__ formatting."""
    if not text:
        return []
    # Split on markers while preserving markers
    parts = []
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__)')
    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            parts.append(('normal', text[last_end:match.start()]))
        content = match.group()[2:-2]
        if match.group().startswith('**'):
            parts.append(('bold', content))
        elif match.group().startswith('*'):
            parts.append(('italic', content))
        elif match.group().startswith('__'):
            parts.append(('underline', content))
        last_end = match.end()
    if last_end < len(text):
        parts.append(('normal', text[last_end:]))
    return parts if parts else [('normal', text)]


def style_run(run, style):
    """Apply character style to a run."""
    if style == 'bold':
        run.bold = True
    elif style == 'italic':
        run.italic = True
    elif style == 'underline':
        run.underline = True


def add_table(doc, table_data):
    """Add a table from CSV-like string."""
    rows = table_data.strip().split('|')
    if not rows:
        return
    cols = [c.strip() for c in rows[0].split(',')]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, col in enumerate(cols):
        cell = hdr.cells[i]
        cell.text = col.strip()
        cell.paragraphs[0].runs[0].bold = True
    # Data rows
    for row_str in rows[1:]:
        cells = [c.strip() for c in row_str.split(',')]
        row = table.add_row()
        for i, cell_text in enumerate(cells):
            if i < len(row.cells):
                row.cells[i].text = cell_text


def add_image(doc, image_spec):
    """
    Add an image. Spec format: "path,width,alignment"
    width in inches, alignment: left/center/right
    """
    parts = [p.strip() for p in image_spec.split(',')]
    path = parts[0]
    width = float(parts[1]) if len(parts) > 1 else 4.0
    align = parts[2] if len(parts) > 2 else 'center'

    if not Path(path).exists():
        print(f"  WARNING: Image not found: {path}")
        return

    paragraph = doc.add_paragraph()
    if align == 'left':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    run.add_picture(path, width=Inches(width))


def add_hyperlink_block(doc, spec):
    """Add a hyperlink block. Spec: 'text|url'"""
    parts = spec.split('|')
    if len(parts) < 2:
        return
    text, url = parts[0].strip(), parts[1].strip()
    paragraph = doc.add_paragraph()
    add_hyperlink(paragraph, text, url)


def process_content_arg(content_str):
    """Parse content string into structured blocks."""
    blocks = []
    raw_blocks = content_str.split('||')
    for raw in raw_blocks:
        raw = raw.strip()
        if not raw:
            continue
        # Detect block type from prefix
        if raw.startswith('Heading 1|') or raw.startswith('H1|'):
            _, *parts = raw.split('|')
            blocks.append(('h1', '|'.join(parts)))
        elif raw.startswith('Heading 2|') or raw.startswith('H2|'):
            _, *parts = raw.split('|')
            blocks.append(('h2', '|'.join(parts)))
        elif raw.startswith('Heading 3|') or raw.startswith('H3|'):
            _, *parts = raw.split('|')
            blocks.append(('h3', '|'.join(parts)))
        elif raw.startswith('Bold|'):
            _, *parts = raw.split('|')
            blocks.append(('bold', '|'.join(parts)))
        elif raw.startswith('Bullet|'):
            items = raw.split('|')[1:]
            blocks.append(('bullet', items))
        elif raw.startswith('Number|'):
            items = raw.split('|')[1:]
            blocks.append(('number', items))
        elif raw.startswith('Roman|'):
            items = raw.split('|')[1:]
            blocks.append(('roman', items))
        elif raw.startswith('Quote|'):
            _, *parts = raw.split('|')
            blocks.append(('quote', '|'.join(parts)))
        elif raw.startswith('HR') and raw == 'HR':
            blocks.append(('hr', ''))
        elif raw.startswith('Image|'):
            parts = raw.split('|')[1:]
            blocks.append(('image', ','.join(parts)))
        elif raw.startswith('Link|'):
            parts = raw.split('|')[1:]
            blocks.append(('hyperlink', '|'.join(parts)))
        elif raw.startswith('Table|'):
            table_str = '|'.join(raw.split('|')[1:])
            blocks.append(('table', table_str))
        else:
            blocks.append(('para', raw))
    return blocks


def build_doc_from_content(doc, blocks, use_styles=True):
    """Build the document from parsed blocks."""
    for block_type, block_data in blocks:
        if block_type == 'h1':
            if use_styles:
                p = doc.add_heading(block_data, level=1)
            else:
                p = doc.add_paragraph(block_data)
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(16)
        elif block_type == 'h2':
            if use_styles:
                p = doc.add_heading(block_data, level=2)
            else:
                p = doc.add_paragraph(block_data)
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(14)
        elif block_type == 'h3':
            if use_styles:
                doc.add_heading(block_data, level=3)
            else:
                p = doc.add_paragraph(block_data)
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(12)
        elif block_type == 'bold':
            p = doc.add_paragraph()
            parts = parse_inline(block_data)
            for style, text in parts:
                run = p.add_run(text)
                style_run(run, style)
        elif block_type == 'para':
            p = doc.add_paragraph()
            parts = parse_inline(block_data)
            for style, text in parts:
                run = p.add_run(text)
                style_run(run, style)
        elif block_type == 'bullet':
            for item in block_data:
                p = doc.add_paragraph(item, style='List Bullet')
                parts = parse_inline(item)
                if len(parts) > 1 or parts[0][0] != 'normal':
                    p.runs[0].text = ''
                    for style, text in parts:
                        run = p.add_run(text)
                        style_run(run, style)
        elif block_type == 'number':
            for item in block_data:
                p = doc.add_paragraph(item, style='List Number')
                parts = parse_inline(item)
                if len(parts) > 1 or parts[0][0] != 'normal':
                    p.runs[0].text = ''
                    for style, text in parts:
                        run = p.add_run(text)
                        style_run(run, style)
        elif block_type == 'roman':
            for i, item in enumerate(block_data):
                roman_numeral = ['i', 'ii', 'iii', 'iv', 'v', 'vi', 'vii', 'viii', 'ix', 'x'][min(i, 9)]
                p = doc.add_paragraph(style='List Number')
                run = p.add_run(f"{roman_numeral}. ")
                run.bold = True
                parts = parse_inline(item)
                for style, text in parts:
                    run = p.add_run(text)
                    style_run(run, style)
        elif block_type == 'quote':
            p = doc.add_paragraph(block_data, style='Quote')
        elif block_type == 'hr':
            p = doc.add_paragraph()
            p_fmt = p.paragraph_format
            p_fmt.space_before = Pt(6)
            p_fmt.space_after = Pt(6)
            run = p.add_run()
            run.add_break()
        elif block_type == 'image':
            add_image(doc, block_data)
        elif block_type == 'hyperlink':
            add_hyperlink_block(doc, block_data)
        elif block_type == 'table':
            add_table(doc, block_data)


def load_from_markdown(filepath):
    """Load content from a markdown file."""
    from markdown_it import MarkdownIt
    md = MarkdownIt()
    content = Path(filepath).read_text(encoding='utf-8')
    tokens = md.parse(content)
    blocks = []
    current_para = []
    in_bullet = False

    def flush_para():
        nonlocal current_para
        if current_para:
            blocks.append(('para', ' '.join(current_para)))
            current_para = []

    for tok in tokens:
        if tok.type == 'heading_open':
            flush_para()
            level = int(tok.tag[1])
            blocks.append(('heading', level))
        elif tok.type == 'heading_close':
            if blocks and blocks[-1][0] == 'heading':
                level, text = blocks.pop()
                blocks.append((f'h{level}', text))
        elif tok.type == 'inline' and blocks and blocks[-1][0] == 'heading':
            blocks[-1] = (blocks[-1][0], tok.content)
        elif tok.type == 'paragraph_open':
            flush_para()
            in_bullet = False
        elif tok.type == 'paragraph_close':
            flush_para()
        elif tok.type == 'bullet_list_open':
            in_bullet = True
        elif tok.type == 'bullet_list_close':
            in_bullet = False
        elif tok.type == 'list_item_open':
            current_bullet = []
        elif tok.type == 'list_item_close':
            if current_bullet:
                blocks.append(('bullet', current_bullet))
        elif tok.type == 'inline' and in_bullet:
            current_bullet.append(tok.content)
        elif tok.type == 'inline' and not in_bullet:
            current_para.append(tok.content)
        elif tok.type == 'hr':
            blocks.append(('hr', ''))

    return blocks


def build_doc_from_markdown(doc, blocks):
    """Build docx from markdown-parsed blocks."""
    for block in blocks:
        if block[0] == 'h1':
            doc.add_heading(block[1], level=1)
        elif block[0] == 'h2':
            doc.add_heading(block[1], level=2)
        elif block[0] == 'h3':
            doc.add_heading(block[1], level=3)
        elif block[0] == 'para':
            doc.add_paragraph(block[1])
        elif block[0] == 'bullet':
            for item in block[1]:
                doc.add_paragraph(item, style='List Bullet')
        elif block[0] == 'hr':
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break()


def main():
    parser = argparse.ArgumentParser(description='Create Word documents')
    parser.add_argument('--output', '-o', required=True, help='Output .docx path')
    parser.add_argument('--title', '-t', help='Document title')
    parser.add_argument('--author', help='Author name')
    parser.add_argument('--content', '-c', help='Pipe-delimited content blocks')
    parser.add_argument('--file', '-f', help='Load content from markdown file')
    parser.add_argument('--table', action='append', help='Add table: "Header1,Header2|Row1C1,R1C2"')
    parser.add_argument('--image', action='append', help='Add image: "path,width,align"')
    parser.add_argument('--hyperlink', action='append', help='Add hyperlink: "text|url"')
    parser.add_argument('--stylesheet', action='store_true', help='Use Word built-in styles')
    parser.add_argument('--no-inline', action='store_true', help='Disable inline formatting')
    args = parser.parse_args()

    doc = Document()

    if args.title:
        doc.core_properties.title = args.title
    if args.author:
        doc.core_properties.author = args.author

    use_styles = args.stylesheet

    if args.file:
        blocks = load_from_markdown(args.file)
        build_doc_from_markdown(doc, blocks)
    elif args.content:
        blocks = process_content_arg(args.content)
        build_doc_from_content(doc, blocks, use_styles=use_styles)

    if args.table:
        for t in args.table:
            add_table(doc, t)

    if args.image:
        for img in args.image:
            add_image(doc, img)

    if args.hyperlink:
        for hl in args.hyperlink:
            add_hyperlink_block(doc, hl)

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"Created: {output} ({output.stat().st_size:,} bytes)")


if __name__ == '__main__':
    main()
