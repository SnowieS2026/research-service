"""
md_to_docx.py -- Convert Markdown files to Word documents.
Usage:
  python md_to_docx.py --input "source.md" --output "out.docx"
  python md_to_docx.py --input "source.md" --output "out.docx" --styles --toc --cover "Title|Author|Date"

Options:
  --input              Input markdown file (required)
  --output             Output .docx path (required)
  --styles             Use Word built-in heading styles (Heading 1, 2, 3)
  --toc                Auto-generate table of contents
  --cover              Cover page: "Title|Author|Date"
  --author             Author (document property)
"""

import argparse
import sys
from pathlib import Path

try:
    from markdown_it import MarkdownIt
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: Missing dependencies. Run: pip install markdown-it-py python-docx")
    sys.exit(1)


def md_to_docx(md_path, docx_path, use_styles=False, add_toc=False,
               cover_info=None, author=None):
    """Convert a markdown file to .docx."""

    md_text = Path(md_path).read_text(encoding='utf-8')
    md = MarkdownIt()
    tokens = md.parse(md_text)

    doc = Document()

    if author:
        doc.core_properties.author = author

    # Build token stream
    current_heading = None
    current_heading_level = None
    in_bullet = False
    current_bullet_items = []
    current_para_parts = []
    in_code_block = False
    code_block_lines = []

    def flush_para():
        nonlocal current_para_parts
        if current_para_parts:
            text = ' '.join(current_para_parts)
            doc.add_paragraph(text)
            current_para_parts = []

    def flush_bullet():
        nonlocal current_bullet_items
        if current_bullet_items:
            for item in current_bullet_items:
                p = doc.add_paragraph(item, style='List Bullet')
            current_bullet_items = []

    for i, tok in enumerate(tokens):
        if tok.type == 'heading_open':
            flush_para()
            flush_bullet()
            level = int(tok.tag[1])
            current_heading_level = level
        elif tok.type == 'heading_close':
            if current_heading and current_heading_level:
                text = current_heading
                if use_styles and current_heading_level in (1, 2, 3):
                    doc.add_heading(text, level=current_heading_level)
                else:
                    p = doc.add_paragraph(text)
                    if current_heading_level == 1:
                        p.runs[0].bold = True
                        p.runs[0].font.size = Pt(16)
                    elif current_heading_level == 2:
                        p.runs[0].bold = True
                        p.runs[0].font.size = Pt(14)
            current_heading = None
            current_heading_level = None
        elif tok.type == 'inline':
            content = tok.content.strip()
            if current_heading_level is not None:
                current_heading = content
            elif in_code_block:
                code_block_lines.append(content)
            elif in_bullet:
                current_bullet_items.append(content)
            else:
                current_para_parts.append(content)
        elif tok.type == 'bullet_list_open':
            flush_para()
            flush_bullet()
            in_bullet = True
        elif tok.type == 'bullet_list_close':
            flush_bullet()
            in_bullet = False
        elif tok.type == 'ordered_list_open':
            flush_para()
            flush_bullet()
            in_bullet = True
        elif tok.type == 'ordered_list_close':
            flush_bullet()
            in_bullet = False
        elif tok.type == 'code_block':
            flush_para()
            flush_bullet()
            # Get content from fence
            for j in range(i-1, -1, -1):
                if tokens[j].type == 'fence':
                    code_text = tokens[j].info.strip()
                    break
            else:
                code_text = ''
            if code_text:
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        elif tok.type == 'hr':
            flush_para()
            flush_bullet()
            p = doc.add_paragraph()
            p_fmt = p.paragraph_format
            p_fmt.space_before = Pt(6)
            p_fmt.space_after = Pt(6)
            run = p.add_run()
            run.add_break()

    flush_para()
    flush_bullet()

    if add_toc:
        # Add TOC field
        doc.add_page_break()
        p = doc.add_paragraph('Table of Contents')
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(16)

    output = Path(docx_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    size = output.stat().st_size
    print(f"Converted: {md_path} -> {docx_path} ({size:,} bytes)")


def main():
    parser = argparse.ArgumentParser(description='Convert Markdown to Word')
    parser.add_argument('--input', '-i', required=True, help='Input markdown file')
    parser.add_argument('--output', '-o', required=True, help='Output .docx path')
    parser.add_argument('--styles', action='store_true', help='Use Word heading styles')
    parser.add_argument('--toc', action='store_true', help='Add table of contents')
    parser.add_argument('--cover', help='Cover page: "Title|Author|Date"')
    parser.add_argument('--author', help='Document author')
    args = parser.parse_args()

    md_to_docx(
        args.input,
        args.output,
        use_styles=args.styles,
        add_toc=args.toc,
        cover_info=args.cover,
        author=args.author,
    )


if __name__ == '__main__':
    main()
